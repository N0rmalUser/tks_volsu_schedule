import logging
import re
from datetime import time
from pathlib import Path

from docx import Document
from docx.table import _Row

from app.core.constants import GROUPS_SCHEDULE_PATH
from app.schemas.enums import WeekType
from app.schemas.schedule import LessonTime, ScheduleRow


DAYS_OF_WEEK = {
    "понедельник": 1,
    "вторник": 2,
    "среда": 3,
    "четверг": 4,
    "пятница": 5,
    "суббота": 6,
}


LESSONS: tuple[LessonTime, ...] = (
    LessonTime(1, time(8, 30), time(10, 0)),
    LessonTime(2, time(10, 10), time(11, 40)),
    LessonTime(3, time(12, 0), time(13, 30)),
    LessonTime(4, time(13, 40), time(15, 10)),
    LessonTime(5, time(15, 20), time(16, 50)),
    LessonTime(6, time(17, 0), time(18, 30)),
    LessonTime(7, time(18, 40), time(20, 10)),
)

LESSON_BY_START_TIME: dict[time, int] = {lesson.start: lesson.number for lesson in LESSONS}


def _parse_info(text: str) -> dict[str, str | list[str] | list[None] | None] | None:
    """Парсит содержимое ячейки расписания."""

    if not text:
        return None

    raw = text.strip()

    # Убираем "- поток N".
    raw = re.sub(
        r"\s*-\s*поток\s*\d+\s*",
        " ",
        raw,
        flags=re.IGNORECASE,
    )

    raw = re.sub(r"\s+", " ", raw)

    parts = re.split(
        r"(?<=\))\s*,\s*",
        raw,
        maxsplit=1,
    )

    if len(parts) == 1:
        parts = re.split(
            r",\s*",
            raw,
            maxsplit=1,
        )

    subject = parts[0].strip()
    rest = parts[1].strip() if len(parts) > 1 else ""

    classroom = ""

    if rest:
        auditorium = re.search(
            r"Ауд\.?\s*([^,;]+)",
            rest,
            flags=re.IGNORECASE,
        )

        if auditorium:
            classroom = re.sub(
                r"\s*",
                "",
                auditorium.group(1),
            )

            classroom = classroom.replace(
                "Спортивныйзал",
                "Спортзал ",
            )

        rest = re.sub(
            r"Ауд\.?\s*([^,;]+)",
            "",
            rest,
            flags=re.IGNORECASE,
        )

    # Убираем должности преподавателей.
    rest = re.sub(
        r"\b("
        r"доцент|"
        r"старший преподаватель|"
        r"старщий преподаватель|"
        r"ассистент|"
        r"профессор|"
        r"преподаватель"
        r")\b\.?",
        "",
        rest,
        flags=re.IGNORECASE,
    )

    teachers = [teacher for teacher in (value.strip() for value in re.split(r"\s*,\s*", rest)) if teacher]

    return {
        "subject": subject,
        "teachers": teachers or [None],
        "classroom": classroom or None,
    }


def _parse_day_and_lesson(row: _Row) -> tuple[int, int]:
    day_name = re.sub(
        r"\s+",
        "",
        row.cells[0].text.strip(),
    ).lower()

    try:
        day_of_week = DAYS_OF_WEEK[day_name]
    except KeyError as exc:
        raise ValueError(
            f"Неизвестный день недели: {day_name!r}",
        ) from exc

    start = row.cells[1].text.strip().split("-", maxsplit=1)[0]
    start = re.sub(r"\s*", "", start)

    try:
        hour, minute = map(int, start.split(":"))
        start_time = time(hour, minute)
    except ValueError as exc:
        raise ValueError(
            f"Не удалось распознать время пары: {start!r}",
        ) from exc

    try:
        lesson_number = LESSON_BY_START_TIME[start_time]
    except KeyError as exc:
        raise ValueError(
            f"Неизвестное время начала пары: {start!r}",
        ) from exc

    return day_of_week, lesson_number


def _process_group_cells(
    left: str,
    right: str | None = None,
    *,
    single_column: bool,
) -> (
    list[tuple[None, dict[str, str | list[str] | list[None] | None]]]
    | list[tuple[int | None, dict[str, str | list[str]]]]
):
    """Разбирает ячейки группы на подгруппы."""

    left = left.strip() if left else ""
    right = right.strip() if right else ""

    if not left and not right:
        return []

    if single_column:
        info = _parse_info(left)

        return [(None, info)] if info else []

    # Одинаковая запись в обеих колонках = общая пара.
    if left and right and left == right:
        info = _parse_info(left)

        return [(None, info)] if info else []

    result: list[tuple[int | None, dict[str, str | list[str]]]] = []

    for subgroup, text in (
        (1, left),
        (2, right),
    ):
        if not text:
            continue

        info = _parse_info(text)

        if not info:
            continue

        subject = str(info["subject"])

        # Лекция считается общей для обеих подгрупп.
        is_lecture = bool(
            re.search(
                r"\((?:Л|Лекция)\)",
                subject,
                flags=re.IGNORECASE,
            )
            or "Лекция" in subject,
        )

        result.append(
            (None if is_lecture else subgroup, info),
        )

    return result


def _extract_groups(
    rows: list[_Row],
    filename: str,
) -> list[tuple[str, int, int | None, bool]]:
    """Определяет группы и соответствующие им колонки."""

    header = [cell.text.strip() for cell in rows[0].cells]

    groups: list[tuple[str, int, int | None, bool]] = []

    column = 2

    while column < len(header):
        group_name = header[column]

        if not group_name:
            column += 1
            continue

        # Две соседние колонки принадлежат одной группе.
        if column + 1 < len(header) and header[column + 1] == group_name:
            groups.append(
                (
                    group_name,
                    column,
                    column + 1,
                    False,
                ),
            )
            column += 2
            continue

        # Одна колонка = одна группа.
        groups.append(
            (
                group_name,
                column,
                None,
                True,
            ),
        )
        column += 1

    if not groups:
        group_name = Path(filename).stem

        if len(rows[0].cells) == 3:
            groups = [
                (group_name, 2, None, True),
            ]
        elif len(rows[0].cells) >= 4:
            groups = [
                (group_name, 2, 3, False),
            ]

    return groups


def _build_rows_for_cell(
    *,
    group_name: str,
    left: str,
    right: str,
    single_column: bool,
    day_of_week: int,
    lesson_number: int,
    week_type: WeekType,
) -> list[ScheduleRow]:
    entries = _process_group_cells(
        left,
        right,
        single_column=single_column,
    )

    result: list[ScheduleRow] = []

    for subgroup, info in entries:
        if not info:
            continue

        subject = info["subject"]

        if not subject:
            continue

        teachers = info["teachers"]
        classroom = info["classroom"]

        # Один преподаватель - одна запись.
        # Если в ячейке несколько преподавателей,
        # создаём несколько ScheduleRow.

        if not teachers:
            teachers = [None]

        result.extend(
            ScheduleRow(
                group=group_name,
                teacher=teacher,
                subject=subject,
                room=classroom,
                day_of_week=day_of_week,
                lesson_number=lesson_number,
                week_type=week_type,
                subgroup=subgroup,
            )
            for teacher in teachers
        )

    return result


def parse_university_schedule() -> list[ScheduleRow]:
    """Парсит все DOCX-файлы расписания."""

    result: list[ScheduleRow] = []
    files = sorted(GROUPS_SCHEDULE_PATH.glob("*.docx"))

    for file_path in files:
        document = Document(str(file_path))

        if not document.tables:
            continue

        rows = document.tables[0].rows

        if not rows:
            continue

        groups = _extract_groups(
            rows,
            file_path.name,
        )

        i = 1

        while i < len(rows):
            row = rows[i]
            day_of_week, lesson_number = _parse_day_and_lesson(row)

            # Если следующая строка имеет тот же день
            # и номер пары — это числитель/знаменатель.
            pair = None

            if i + 1 < len(rows):
                next_day, next_lesson = _parse_day_and_lesson(
                    rows[i + 1],
                )

                if next_day == day_of_week and next_lesson == lesson_number:
                    pair = rows[i + 1]

            for (
                group_name,
                col1,
                col2,
                single_column,
            ) in groups:
                if col1 >= len(row.cells):
                    continue

                left = row.cells[col1].text

                right = row.cells[col2].text if col2 is not None and col2 < len(row.cells) else ""

                if pair is not None:
                    # Числитель.
                    result.extend(
                        _build_rows_for_cell(
                            group_name=group_name,
                            left=left,
                            right=right,
                            single_column=single_column,
                            day_of_week=day_of_week,
                            lesson_number=lesson_number,
                            week_type=WeekType.ODD,
                        ),
                    )

                    # Знаменатель.
                    pair_left = pair.cells[col1].text

                    pair_right = pair.cells[col2].text if col2 is not None and col2 < len(pair.cells) else ""

                    result.extend(
                        _build_rows_for_cell(
                            group_name=group_name,
                            left=pair_left,
                            right=pair_right,
                            single_column=single_column,
                            day_of_week=day_of_week,
                            lesson_number=lesson_number,
                            week_type=WeekType.EVEN,
                        ),
                    )

                else:
                    # Одна запись действует каждую неделю.
                    result.extend(
                        _build_rows_for_cell(
                            group_name=group_name,
                            left=left,
                            right=right,
                            single_column=single_column,
                            day_of_week=day_of_week,
                            lesson_number=lesson_number,
                            week_type=WeekType.EVERY,
                        ),
                    )

            i += 2 if pair is not None else 1
    logging.info("Расписания университета успешно сохранены в базу данных.")
    return result
